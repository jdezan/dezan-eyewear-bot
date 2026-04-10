"""
Processa PDFs e DOCX da pasta data/ e gera o vectorstore FAISS.

Uso:
  python ingest.py

Este script deve ser executado:
  1. Na primeira vez que subir o bot
  2. Sempre que adicionar/atualizar catálogos em data/
"""

import os
import sys
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain_community.vectorstores import FAISS
from langchain.schema import Document
import config


def _build_embeddings():
    """Cria objeto de embeddings conforme EMBEDDING_PROVIDER."""
    if config.EMBEDDING_PROVIDER == "huggingface":
        from langchain_community.embeddings import HuggingFaceEmbeddings
        print(f"  Usando embeddings HuggingFace: {config.EMBEDDING_MODEL}")
        return HuggingFaceEmbeddings(
            model_name=config.EMBEDDING_MODEL,
            encode_kwargs={"normalize_embeddings": True},
        )
    from langchain_openai import OpenAIEmbeddings
    print("  Usando embeddings OpenAI")
    return OpenAIEmbeddings(api_key=config.OPENAI_API_KEY)


def load_docx(path: str) -> list[Document]:
    """Carrega um .docx e retorna lista de Documents (parágrafos + tabelas)."""
    from docx import Document as DocxDoc

    doc = DocxDoc(path)
    text = "\n".join(p.text for p in doc.paragraphs if p.text.strip())

    # Extrai conteúdo de tabelas também
    for table in doc.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells if c.text.strip()]
            if cells:
                text += "\n" + " | ".join(cells)

    return [Document(page_content=text, metadata={"source": os.path.basename(path)})]


def load_pdf(path: str) -> list[Document]:
    """Carrega um .pdf e retorna lista de Documents."""
    from langchain_community.document_loaders import PyPDFLoader
    return PyPDFLoader(path).load()


def ingest():
    # Coleta arquivos
    files = [
        os.path.join(config.DATA_DIR, f)
        for f in sorted(os.listdir(config.DATA_DIR))
        if f.lower().endswith((".pdf", ".docx"))
    ]

    if not files:
        print(f"[ERRO] Nenhum PDF/DOCX encontrado em '{config.DATA_DIR}/'")
        sys.exit(1)

    print(f"\n📂 Processando {len(files)} arquivo(s) de '{config.DATA_DIR}/'...")

    docs: list[Document] = []
    for fpath in files:
        name = os.path.basename(fpath)
        print(f"  → {name}")
        try:
            if fpath.lower().endswith(".docx"):
                docs.extend(load_docx(fpath))
            else:
                docs.extend(load_pdf(fpath))
        except Exception as e:
            print(f"    [AVISO] Erro ao ler {name}: {e}")

    print(f"\n✅ {len(docs)} documento(s) carregado(s)")

    # Divide em chunks
    splitter = RecursiveCharacterTextSplitter(
        chunk_size=config.CHUNK_SIZE,
        chunk_overlap=config.CHUNK_OVERLAP,
        separators=["\n\n", "\n", ". ", " "],
    )
    chunks = splitter.split_documents(docs)
    print(f"✂️  {len(chunks)} chunks gerados (tamanho={config.CHUNK_SIZE}, overlap={config.CHUNK_OVERLAP})")

    # Gera embeddings e salva
    print("\n🔢 Gerando embeddings e salvando vectorstore...")
    embeddings = _build_embeddings()
    vectorstore = FAISS.from_documents(chunks, embeddings)
    vectorstore.save_local(config.VECTORSTORE_DIR)

    print(f"\n🎉 Vectorstore salvo em '{config.VECTORSTORE_DIR}/'")
    print("   Próximo passo: python api.py\n")


if __name__ == "__main__":
    ingest()
